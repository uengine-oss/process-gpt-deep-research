import json
import logging
import os
import re
import tempfile
import uuid
from concurrent.futures import ThreadPoolExecutor
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple
from urllib.parse import quote, unquote, urlparse

import requests
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph

from ..db import get_db_client
from .images import generate_image
from .llm import chat_text



logger = logging.getLogger("docx-template")

PLACEHOLDER_PATTERN = re.compile(r"\[[^\[\]]+?\]")
HEADING_NUMBER_PATTERN = re.compile(r"^\s*\d+(\.\d+)*\s+")
PARA_RANGE_PATTERN = re.compile(r"(\d+)\s*[~\-]\s*(\d+)\s*문단")
PARA_COUNT_PATTERN = re.compile(r"(\d+)\s*문단")
CHAR_COUNT_PATTERN = re.compile(r"(\d+)\s*자")
DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
STORAGE_BUCKET = "deep_research_files"


def _extract_public_url(response: Any) -> Optional[str]:
    if not response:
        return None
    if isinstance(response, dict):
        if response.get("publicUrl"):
            return response.get("publicUrl")
        if response.get("public_url"):
            return response.get("public_url")
        data = response.get("data")
        if isinstance(data, dict) and data.get("publicUrl"):
            return data.get("publicUrl")
    return None


def _extract_storage_path_from_public_url(url: str) -> Optional[str]:
    if not url:
        return None
    parsed = urlparse(url)
    marker = f"/storage/v1/object/public/{STORAGE_BUCKET}/"
    if marker not in parsed.path:
        return None
    path_part = parsed.path.split(marker, 1)[-1]
    return unquote(path_part).lstrip("/")


def _download_docx(url: str) -> Path:
    resp = requests.get(url, timeout=30)
    resp.raise_for_status()
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp.write(resp.content)
    tmp.flush()
    tmp.close()
    return Path(tmp.name)


def _has_outline_level(para) -> bool:
    try:
        ppr = para._p.pPr
        if ppr is None or ppr.outlineLvl is None:
            return False
        val = int(ppr.outlineLvl.val)
        return val <= 1
    except Exception:
        return False


def _is_heading(para) -> bool:
    style_name = para.style.name if para.style else ""
    style_lower = str(style_name or "").lower()
    if style_lower.startswith("heading"):
        return True
    # localized/common heading style names
    if "제목" in style_name or "표제" in style_name:
        return True
    if _has_outline_level(para):
        return True
    return False


def _heading_level(para) -> str:
    style_name = para.style.name if para.style else ""
    return str(style_name or "").strip()


def _heading_depth(para, para_text: str) -> Optional[int]:
    text = (para_text or "").strip()
    number_match = re.match(r"^\s*(\d+(?:\.\d+)*)\s+", text)
    if number_match:
        return len(number_match.group(1).split("."))
    style_name = para.style.name if para.style else ""
    style_match = re.search(r"(\d+)", str(style_name or ""))
    if style_match:
        try:
            return int(style_match.group(1))
        except Exception:
            return None
    return None


def _set_paragraph_text(para, text: str) -> None:
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(text)


def _insert_paragraph_after(paragraph, text: str = ""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph.__class__(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def _iter_block_items(parent):
    for child in parent.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _looks_like_heading(para) -> bool:
    if not para.text:
        return False
    return bool(HEADING_NUMBER_PATTERN.match(para.text.strip()))


def _normalize_paragraph_text(text: str) -> str:
    if text is None:
        return ""
    normalized = str(text)
    normalized = normalized.replace("\u00a0", " ")
    normalized = re.sub(r"[\u200b\u200c\u200d\ufeff]", "", normalized)
    normalized = re.sub(r"\s+", " ", normalized).strip()
    return normalized


def _preview_text(text: str, limit: int = 120) -> str:
    if text is None:
        return ""
    raw = str(text).replace("\n", " ").replace("\r", " ")
    return (raw[:limit] + "...") if len(raw) > limit else raw


def _extract_guidance(text: str) -> List[str]:
    if not text:
        return []
    guidance = []
    # bracketed hints
    for match in PLACEHOLDER_PATTERN.findall(text):
        cleaned = match.strip("[]").strip()
        if cleaned:
            guidance.append(cleaned)
    # inline hints without brackets
    if PARA_RANGE_PATTERN.search(text) or PARA_COUNT_PATTERN.search(text) or CHAR_COUNT_PATTERN.search(text):
        guidance.append(text.strip())
    return guidance


def _merge_length_hints(target: Dict[str, Any], guidance_texts: List[str]) -> None:
    for text in guidance_texts:
        range_match = PARA_RANGE_PATTERN.search(text)
        if range_match:
            min_p = int(range_match.group(1))
            max_p = int(range_match.group(2))
            target["min_paragraphs"] = min_p
            target["max_paragraphs"] = max_p
            continue
        count_match = PARA_COUNT_PATTERN.search(text)
        if count_match:
            count = int(count_match.group(1))
            target.setdefault("min_paragraphs", count)
            target.setdefault("max_paragraphs", count)
        char_match = CHAR_COUNT_PATTERN.search(text)
        if char_match:
            target["max_chars"] = int(char_match.group(1))


def _remove_paragraph(para) -> None:
    try:
        p = para._element
        p.getparent().remove(p)
        p._p = p._element = None
    except Exception:
        pass


def _remove_table(table) -> None:
    try:
        tbl = table._element
        tbl.getparent().remove(tbl)
        tbl._tbl = tbl._element = None
    except Exception:
        pass


def _collect_placeholders_from_paragraphs(paragraphs: Iterable, collector: set) -> None:
    for para in paragraphs:
        if not para.text:
            continue
        collector.update(PLACEHOLDER_PATTERN.findall(para.text))


def _collect_placeholders_from_tables(tables: Iterable, collector: set) -> None:
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                _collect_placeholders_from_paragraphs(cell.paragraphs, collector)


def extract_placeholders(doc: Document) -> List[str]:
    placeholders: set = set()
    _collect_placeholders_from_paragraphs(doc.paragraphs, placeholders)
    _collect_placeholders_from_tables(doc.tables, placeholders)
    for section in doc.sections:
        for part in [
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
        ]:
            if part is None:
                continue
            _collect_placeholders_from_paragraphs(part.paragraphs, placeholders)
            _collect_placeholders_from_tables(part.tables, placeholders)
    return sorted(placeholders)


def extract_template_schema(doc: Document) -> Dict[str, Any]:
    sections: List[Dict[str, Any]] = []
    tables: List[Dict[str, Any]] = []
    current = None
    index_map = {para._p: idx for idx, para in enumerate(doc.paragraphs)}
    max_template_excerpt_chars = 800
    preface_logged = 0
    cover_paragraphs: List[Dict[str, Any]] = []
    cover_tables: List[Dict[str, Any]] = []
    cover_active = True

    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            para = block
            raw_text = para.text or ""
            para_text = _normalize_paragraph_text(raw_text)
            is_heading = _is_heading(para) or _looks_like_heading(para)
            if is_heading:
                cover_active = False
                section_id = f"section_{len(sections)+1}"
                current = {
                    "id": section_id,
                    "title": para_text,
                    "level": _heading_level(para) or "Heading",
                    "depth": _heading_depth(para, para_text),
                    "heading_index": index_map.get(para._p),
                    "paragraph_indices": [],
                    "template_texts": [],
                    "optional": False,
                }
                sections.append(current)
                logger.debug(
                    "DOCX 섹션 헤딩 감지: id=%s title=%s raw=%s normalized=%s",
                    section_id,
                    _preview_text(para_text),
                    _preview_text(raw_text),
                    _preview_text(para_text),
                )
                continue
            if not sections and para_text and preface_logged < 6:
                style_name = para.style.name if para.style else ""
                logger.debug(
                    "DOCX 표지 후보 문단 감지: idx=%s style=%s raw=%s normalized=%s",
                    index_map.get(para._p),
                    style_name,
                    _preview_text(raw_text),
                    _preview_text(para_text),
                )
                preface_logged += 1
            if cover_active and para_text:
                cover_paragraphs.append(
                    {
                        "index": index_map.get(para._p),
                        "text": para_text,
                        "style": para.style.name if para.style else "",
                    }
                )
            if current is not None:
                idx = index_map.get(para._p)
                if idx is not None and para_text:
                    current["paragraph_indices"].append(idx)
                    current["template_texts"].append(para_text)
                    logger.debug(
                        "DOCX 섹션 본문 문단 감지: id=%s title=%s idx=%s raw=%s normalized=%s",
                        current.get("id"),
                        _preview_text(current.get("title") or ""),
                        idx,
                        _preview_text(raw_text),
                        _preview_text(para_text),
                    )
                    guidance = _extract_guidance(raw_text)
                    if guidance:
                        current.setdefault("guidance", [])
                        current["guidance"].extend(guidance)
                        _merge_length_hints(current, guidance)
        elif isinstance(block, Table):
            table = block
            headers = []
            if table.rows:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
            row_samples = []
            for row in table.rows[:6]:
                row_samples.append([cell.text.strip() for cell in row.cells])
            if cover_active:
                cover_tables.append(
                    {
                        "index": len(tables),
                        "rows": row_samples,
                    }
                )
            key_value_no_header = False
            if len(table.columns) == 2 and row_samples:
                label_rows = 0
                for r in row_samples:
                    if not isinstance(r, list) or len(r) < 2:
                        continue
                    left = _normalize_paragraph_text(r[0])
                    right = _normalize_paragraph_text(r[1])
                    if not left:
                        continue
                    if (not right) or right == left:
                        label_rows += 1
                if label_rows >= max(2, len(row_samples) - 1):
                    key_value_no_header = True
            header_is_data = False
            if headers:
                for cell_text in headers:
                    if PLACEHOLDER_PATTERN.search(cell_text):
                        header_is_data = True
                        break
                    if re.search(r"YYYY|MM|DD|\\d{4}", cell_text):
                        header_is_data = True
                        break
            else:
                header_is_data = True
            if key_value_no_header:
                header_is_data = True
            table_id = f"table_{len(tables)+1}"
            tables.append(
                {
                    "id": table_id,
                    "index": len(tables),
                    "headers": headers,
                    "columns": len(table.columns),
                    "rows": len(table.rows),
                    "section_id": current.get("id") if current else None,
                    "section_title": current.get("title") if current else None,
                    "row_samples": row_samples,
                    "header_is_data": header_is_data,
                    "key_value_no_header": key_value_no_header,
                }
            )

    section_ids_with_tables = {t.get("section_id") for t in tables if t.get("section_id")}
    for sec in sections:
        sec["has_tables"] = sec.get("id") in section_ids_with_tables
        template_texts = sec.get("template_texts") or []
        if not isinstance(template_texts, list) or not template_texts:
            continue
        excerpt = "\n\n".join(str(t) for t in template_texts if t).strip()
        if excerpt and len(excerpt) > max_template_excerpt_chars:
            excerpt = excerpt[:max_template_excerpt_chars].rstrip()
        sec["template_excerpt"] = excerpt

    for i, sec in enumerate(sections):
        current_depth = sec.get("depth")
        has_children: Optional[bool] = None
        if isinstance(current_depth, int):
            has_children = False
            next_depth = None
            if i + 1 < len(sections):
                next_depth = sections[i + 1].get("depth")
            if isinstance(next_depth, int):
                has_children = next_depth > current_depth
            else:
                has_children = None
        sec["has_children"] = has_children
    return {
        "sections": sections,
        "tables": tables,
        "cover": {"paragraphs": cover_paragraphs, "tables": cover_tables},
    }


def summarize_template_schema(schema: Dict[str, Any], max_chars: int = 2000) -> str:
    lines: List[str] = []
    for sec in schema.get("sections", []):
        title = sec.get("title") or sec.get("id")
        optional = " optional" if sec.get("optional") else ""
        guidance = sec.get("guidance") or []
        guidance_text = f" guidance={'; '.join(guidance[:2])}" if guidance else ""
        lines.append(f"- SECTION {sec.get('id')}: {title}{optional}{guidance_text}")
    for tbl in schema.get("tables", []):
        headers = ", ".join([h for h in (tbl.get("headers") or []) if h])
        sec_title = tbl.get("section_title") or ""
        sec_hint = f" section={sec_title}" if sec_title else ""
        lines.append(f"- TABLE {tbl.get('id')}: headers=[{headers}]{sec_hint}")
    summary = "\n".join(lines)
    return summary[:max_chars]


def load_template_schema_summary(template_url: str) -> str:
    if not template_url:
        return ""
    try:
        template_path = _download_docx(template_url)
        doc = Document(str(template_path))
        schema = extract_template_schema(doc)
        logger.debug(
            "템플릿 스키마 요약: sections=%s tables=%s sample_titles=%s",
            len(schema.get("sections") or []),
            len(schema.get("tables") or []),
            [s.get("title") for s in (schema.get("sections") or [])[:5]],
        )
        return summarize_template_schema(schema)
    except Exception as e:
        logger.error("템플릿 스키마 추출 실패: %s", e)
        return ""


def load_template_schema(template_url: str) -> Dict[str, Any]:
    if not template_url:
        return {"sections": [], "tables": []}
    try:
        template_path = _download_docx(template_url)
        doc = Document(str(template_path))
        schema = extract_template_schema(doc)
        logger.debug(
            "템플릿 스키마 로드: sections=%s tables=%s sample_titles=%s",
            len(schema.get("sections") or []),
            len(schema.get("tables") or []),
            [s.get("title") for s in (schema.get("sections") or [])[:5]],
        )
        # 디버깅: paragraph_indices가 빈 섹션 = doc.add_paragraph 대상
        empty_idx = [(s.get("id"), s.get("title")) for s in (schema.get("sections") or []) if not (s.get("paragraph_indices"))]
        if empty_idx:
            logger.debug(
                "템플릿 스키마 섹션(paragraph_indices 없음): %s → 본문 삽입 위치 필요",
                empty_idx,
            )
        return schema
    except Exception as e:
        logger.error("템플릿 스키마 로드 실패: %s", e)
        return {"sections": [], "tables": []}


def _replace_in_paragraph(para, replacements: Dict[str, str]) -> int:
    if not para.runs:
        return 0
    full_text = "".join(run.text for run in para.runs)
    if not any(ph in full_text for ph in replacements):
        return 0

    replaced = 0
    for run in para.runs:
        original = run.text
        for ph, val in replacements.items():
            if ph in run.text:
                run.text = run.text.replace(ph, val)
        if run.text != original:
            replaced += 1

    if replaced == 0:
        new_text = full_text
        for ph, val in replacements.items():
            if ph in new_text:
                new_text = new_text.replace(ph, val)
        if new_text != full_text:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
            replaced += 1
    return replaced


def replace_placeholders(doc: Document, replacements: Dict[str, str]) -> int:
    count = 0

    for para in doc.paragraphs:
        count += _replace_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    count += _replace_in_paragraph(para, replacements)

    for section in doc.sections:
        for part in [
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
        ]:
            if part is None:
                continue
            for para in part.paragraphs:
                count += _replace_in_paragraph(para, replacements)
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            count += _replace_in_paragraph(para, replacements)

    return count


def _extract_json(text: str) -> Optional[Dict[str, Any]]:
    if not text:
        return None
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1 or end <= start:
        return None
    try:
        return json.loads(text[start : end + 1])
    except Exception:
        return None


def build_schema_output(schema: Dict[str, Any], report_markdown: str, query: str) -> Dict[str, Any]:
    sections = schema.get("sections") or []
    tables = schema.get("tables") or []
    if not sections and not tables:
        return {}

    compact_schema = {
        "sections": [
            {
                "id": item.get("id"),
                "title": item.get("title"),
                "level": item.get("level"),
                "optional": item.get("optional"),
            }
            for item in sections
        ],
        "tables": [
            {
                "id": item.get("id"),
                "headers": item.get("headers"),
                "columns": item.get("columns"),
                "section_id": item.get("section_id"),
            }
            for item in tables
        ],
    }

    prompt_sys = "You are a document template filler. Return JSON only."
    prompt_user = (
        "아래 템플릿 스키마에 맞춰 결과를 채워 주세요.\n"
        "- 반드시 JSON만 출력하세요.\n"
        "- sections: 각 section.id에 대해 status와 content를 작성하세요.\n"
        "  - status는 fill | partial | omit 중 하나입니다.\n"
        "  - optional=true 섹션은 자료가 부족하면 omit으로 처리하세요.\n"
        "- tables: 각 table.id에 대해 status와 rows 배열(행의 2차원 배열)을 작성하세요.\n"
        "  - status는 fill | partial | omit 중 하나입니다.\n"
        "- images: 이미지가 유용한 경우에만 배열로 작성하세요. 각 항목은 {section_id, prompt, caption}.\n"
        "- 표의 행 수는 필요에 따라 작성하되, 헤더 개수와 열 수에 맞춰 주세요.\n\n"
        f"[사용자 요청]\n{query}\n\n"
        f"[템플릿 스키마]\n{json.dumps(compact_schema, ensure_ascii=False)}\n\n"
        f"[리서치 결과]\n{report_markdown}\n"
    )
    raw = chat_text(prompt_sys, prompt_user) or ""
    return _extract_json(raw) or {}


def apply_schema_output(
    doc: Document,
    schema: Dict[str, Any],
    output: Dict[str, Any],
    report_id: Optional[str] = None,
) -> None:
    cover_output = output.get("cover") if isinstance(output, dict) else None
    if isinstance(cover_output, dict):
        title_index = cover_output.get("title_index")
        subtitle_index = cover_output.get("subtitle_index")
        title_text = (cover_output.get("title_text") or "").strip()
        subtitle_text = (cover_output.get("subtitle_text") or "").strip()
        if isinstance(title_index, int) and 0 <= title_index < len(doc.paragraphs) and title_text:
            _set_paragraph_text(doc.paragraphs[title_index], title_text)
            logger.debug(
                "apply_schema: cover_title_update index=%s text=%s",
                title_index,
                title_text[:120],
            )
        if (
            isinstance(subtitle_index, int)
            and 0 <= subtitle_index < len(doc.paragraphs)
            and subtitle_text
        ):
            _set_paragraph_text(doc.paragraphs[subtitle_index], subtitle_text)
            logger.debug(
                "apply_schema: cover_subtitle_update index=%s text=%s",
                subtitle_index,
                subtitle_text[:120],
            )

    sections_output = output.get("sections") if isinstance(output, dict) else None
    omitted_section_ids: set = set()
    if isinstance(sections_output, dict):
        section_index = {s.get("id"): s for s in schema.get("sections", [])}
        for sec in schema.get("sections", []):
            new_title = sec.get("mapped_title")
            heading_index = sec.get("heading_index")
            if not new_title:
                continue
            if heading_index is None:
                continue
            if heading_index >= len(doc.paragraphs):
                continue
            _set_paragraph_text(doc.paragraphs[heading_index], str(new_title))
            logger.debug(
                "apply_schema: heading_title_update section=%s heading_index=%s new_title=%s",
                sec.get("id"),
                heading_index,
                str(new_title)[:120],
            )
        for sec_id, content in sections_output.items():
            sec = section_index.get(sec_id)
            if not sec:
                continue
            status = None
            text = ""
            if isinstance(content, dict):
                status = (content.get("status") or "").strip().lower()
                text = str(content.get("content") or "").strip()
            elif isinstance(content, str):
                text = content.strip()
            if status == "omit":
                # Disable optional deletion again; keep headings/sections for user cleanup.
                status = "partial"
            if not text:
                continue
            indices = sec.get("paragraph_indices") or []
            if indices:
                _set_paragraph_text(doc.paragraphs[indices[0]], text)
                for idx in indices[1:]:
                    _set_paragraph_text(doc.paragraphs[idx], "")
            else:
                doc.add_paragraph(text)
                logger.debug(
                    "apply_schema: section=%s paragraph_indices 없음 → 문서 끝에 추가 (content_len=%s)",
                    sec_id,
                    len(text),
                )

    tables_output = output.get("tables") if isinstance(output, dict) else None
    table_status_by_id: Dict[str, str] = {}
    if isinstance(tables_output, dict):
        for tbl_id, tbl_content in tables_output.items():
            if isinstance(tbl_content, dict):
                table_status_by_id[tbl_id] = str(tbl_content.get("status") or "").strip().lower()
    if isinstance(tables_output, dict):
        table_index = {t.get("id"): t for t in schema.get("tables", [])}
        for tbl_id, tbl_content in tables_output.items():
            tbl_meta = table_index.get(tbl_id)
            if not tbl_meta:
                continue
            idx = tbl_meta.get("index")
            if idx is None or idx >= len(doc.tables):
                continue
            table = doc.tables[idx]
            header_is_data = bool(tbl_meta.get("header_is_data"))
            status = None
            rows = []
            headers_override = None
            if isinstance(tbl_content, dict):
                status = (tbl_content.get("status") or "").strip().lower()
                rows = tbl_content.get("rows") or []
                headers_override = tbl_content.get("headers")
            if not isinstance(rows, list):
                continue
            if status == "omit":
                _remove_table(table)
                continue

            header_row = table.rows[0] if table.rows else None
            data_template_row = table.rows[1] if len(table.rows) > 1 else header_row

            rows_for_data = rows
            if not header_is_data:
                # Update header row using rows[0] when provided
                header_values = None
                if isinstance(headers_override, list) and headers_override:
                    header_values = headers_override
                elif rows:
                    header_values = rows[0] if isinstance(rows[0], list) else []
                if header_row is not None and header_values is not None:
                    for i, cell in enumerate(header_row.cells):
                        value = header_values[i] if i < len(header_values) else ""
                        if cell.paragraphs:
                            _set_paragraph_text(cell.paragraphs[0], str(value))
                            for extra in cell.paragraphs[1:]:
                                _set_paragraph_text(extra, "")
                        else:
                            cell.text = str(value)
                    if header_values is headers_override:
                        rows_for_data = rows
                    else:
                        rows_for_data = rows[1:]

            template_row_xml = deepcopy(data_template_row._tr) if data_template_row is not None else None

            while len(table.rows) > (1 if not header_is_data else 1):
                table._tbl.remove(table.rows[1]._tr)

            for row_index, row_data in enumerate(rows_for_data):
                if not isinstance(row_data, list):
                    continue
                if row_index == 0 and header_is_data and header_row is not None:
                    new_row = header_row
                elif template_row_xml is None:
                    table.add_row()
                else:
                    table._tbl.append(deepcopy(template_row_xml))
                if row_index == 0 and header_is_data and header_row is not None:
                    new_row = header_row
                else:
                    new_row = table.rows[-1]
                for i, cell in enumerate(new_row.cells):
                    value = row_data[i] if i < len(row_data) else ""
                    if cell.paragraphs:
                        _set_paragraph_text(cell.paragraphs[0], str(value))
                        for extra in cell.paragraphs[1:]:
                            _set_paragraph_text(extra, "")
                    else:
                        cell.text = str(value)

    if omitted_section_ids:
        tables = schema.get("tables", []) or []
        removable = []
        for tbl in tables:
            tbl_id = tbl.get("id")
            sec_id = tbl.get("section_id")
            status = table_status_by_id.get(tbl_id)
            if sec_id in omitted_section_ids and status not in ("fill", "partial"):
                if isinstance(tbl.get("index"), int):
                    removable.append(tbl.get("index"))
        for idx in sorted(set(removable), reverse=True):
            if idx < len(doc.tables):
                _remove_table(doc.tables[idx])

    images_output = output.get("images") if isinstance(output, dict) else None
    if isinstance(images_output, list):
        section_index = {s.get("id"): s for s in schema.get("sections", [])}
        image_jobs = []
        for index, item in enumerate(images_output, start=1):
            if not isinstance(item, dict):
                continue
            section_id = item.get("section_id")
            prompt = (item.get("prompt") or "").strip()
            caption = (item.get("caption") or "").strip()
            if not section_id or not prompt:
                continue
            if section_id not in section_index:
                continue
            filename = f"image-{index}.png"
            image_jobs.append((section_id, prompt, caption, filename))

        def _render_image(job):
            section_id, prompt, caption, filename = job
            try:
                tmp_dir = Path(tempfile.mkdtemp())
                img_path = tmp_dir / filename
                if generate_image(prompt, img_path):
                    public_url = None
                    if report_id:
                        public_url = _upload_image(img_path, report_id, filename)
                    return (section_id, img_path, caption, public_url)
            except Exception:
                return None
            return None

        if image_jobs:
            results = []
            with ThreadPoolExecutor(max_workers=5) as executor:
                futures = [executor.submit(_render_image, job) for job in image_jobs]
                for future in futures:
                    try:
                        result = future.result()
                        if result:
                            results.append(result)
                    except Exception:
                        continue

            for section_id, img_path, caption, _public_url in results:
                sec = section_index.get(section_id)
                if not sec:
                    continue
                indices = sec.get("paragraph_indices") or []
                if indices:
                    insert_para = doc.paragraphs[indices[-1]]
                    new_para = _insert_paragraph_after(insert_para)
                else:
                    new_para = doc.add_paragraph()
                try:
                    new_para.add_run().add_picture(str(img_path), width=Inches(5.5))
                    if caption:
                        doc.add_paragraph(caption)
                except Exception:
                    continue

def build_placeholder_mapping(placeholders: List[str], report_markdown: str, query: str) -> Dict[str, str]:
    if not placeholders:
        return {}
    prompt_sys = "You are a document template filler. Return JSON only."
    prompt_user = (
        "다음 템플릿 플레이스홀더를 보고, 리서치 결과를 기준으로 각 항목에 들어갈 텍스트를 매핑하세요.\n"
        "- 반드시 JSON 객체만 출력하세요.\n"
        "- 키는 템플릿 플레이스홀더 문자열을 그대로 사용하세요.\n"
        "- 값은 간결한 문장/단락으로 작성하세요.\n\n"
        f"[사용자 요청]\n{query}\n\n"
        f"[플레이스홀더]\n{placeholders}\n\n"
        f"[리서치 결과]\n{report_markdown}\n"
    )
    raw = chat_text(prompt_sys, prompt_user) or ""
    data = _extract_json(raw) or {}

    mapping: Dict[str, str] = {}
    for ph in placeholders:
        val = data.get(ph)
        if val is None and ph.startswith("[") and ph.endswith("]"):
            val = data.get(ph[1:-1])
        if val is None:
            val = ""
        mapping[ph] = str(val)
    return mapping


def _build_output_filename() -> str:
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
    return f"report-{stamp}.docx"


def _upload_docx(file_path: Path, storage_path: str) -> Optional[str]:
    supabase = get_db_client()
    file_bytes = file_path.read_bytes()
    safe_path = storage_path.lstrip("/")
    try:
        resp = supabase.storage.from_(STORAGE_BUCKET).upload(
            safe_path,
            file_bytes,
            {"content-type": DOCX_CONTENT_TYPE, "upsert": "true"},
        )
        if hasattr(resp, "path") and not resp.path:
            logger.error("storage 업로드 실패: 응답 path 없음 %s", resp)
            return None
        public = supabase.storage.from_(STORAGE_BUCKET).get_public_url(safe_path)
        url = _extract_public_url(public)
        if url:
            return url
    except Exception as e:
        logger.error("storage 업로드 실패: %s", e)
        return None
    base_url = (os.getenv("SUPABASE_URL") or "").rstrip("/")
    if base_url:
        return f"{base_url}/storage/v1/object/public/{STORAGE_BUCKET}/{quote(safe_path, safe='/-_.')}"
    return None


def _upload_image(file_path: Path, report_id: str, filename: str) -> Optional[str]:
    supabase = get_db_client()
    file_bytes = file_path.read_bytes()
    safe_name = Path(filename).name
    storage_path = f"deep-research/{report_id}/{safe_name}"
    safe_path = storage_path.lstrip("/")
    try:
        resp = supabase.storage.from_(STORAGE_BUCKET).upload(
            safe_path,
            file_bytes,
            {"content-type": "image/png", "upsert": "true"},
        )
        if hasattr(resp, "path") and not resp.path:
            logger.error("image 업로드 실패: 응답 path 없음 %s", resp)
            return None
        public = supabase.storage.from_(STORAGE_BUCKET).get_public_url(safe_path)
        url = _extract_public_url(public)
        if url:
            logger.debug("image 업로드 완료: %s", url)
            return url
    except Exception as e:
        logger.error("image 업로드 실패: %s", e)
        return None
    base_url = (os.getenv("SUPABASE_URL") or "").rstrip("/")
    if base_url:
        return f"{base_url}/storage/v1/object/public/{STORAGE_BUCKET}/{quote(safe_path, safe='/-_.')}"
    return None


def _save_proc_inst_source(proc_inst_id: str, file_name: str, file_path: str) -> None:
    supabase = get_db_client()
    payload = {
        "id": str(uuid.uuid4()),
        "proc_inst_id": proc_inst_id,
        "file_name": file_name,
        "file_path": file_path,
        "is_process": True,
    }
    supabase.table("proc_inst_source").insert(payload).execute()


def generate_docx_from_template(
    *,
    template_url: str,
    template_name: str,
    report_markdown: str,
    query: str,
    proc_inst_id: str,
    report_id: str,
    output_name: Optional[str] = None,
    output_display_name: Optional[str] = None,
    schema: Optional[Dict[str, Any]] = None,
    schema_output: Optional[Dict[str, Any]] = None,
) -> Optional[Dict[str, str]]:
    safe_template_name = Path(template_name or "template.docx").name
    template_path = _download_docx(template_url)
    doc = Document(str(template_path))

    if not schema:
        schema = extract_template_schema(doc)
    if not schema_output:
        schema_output = build_schema_output(schema, report_markdown, query)
    if schema and schema_output:
        apply_schema_output(doc, schema, schema_output, report_id=report_id)

    output_name = output_name or _build_output_filename()
    output_display_name = output_display_name or output_name
    output_path = Path(tempfile.mkdtemp()) / output_name
    doc.save(str(output_path))

    storage_path = f"deep-research/{report_id}/{output_name}"

    public_url = _upload_docx(output_path, storage_path)
    if not public_url:
        return None

    _save_proc_inst_source(proc_inst_id, output_display_name, public_url)
    return {
        "file_name": output_display_name,
        "file_path": public_url,
        "storage_path": storage_path,
    }

